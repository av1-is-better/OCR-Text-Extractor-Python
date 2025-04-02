import os
import shutil
from colorama import Fore, Style, init
from PIL import Image
import progressbar
import pytesseract
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_BREAK
import time
import sys
from dotenv import load_dotenv


def print_function(status):
    if status == "intro":
        os.system('cls' if os.name == 'nt' else 'clear')
        print(Fore.LIGHTRED_EX + "===================================\n"+Style.BRIGHT + Fore.YELLOW +"            Python O.C.R            \n"+ Style.RESET_ALL + Fore.LIGHTRED_EX +"========"+Fore.LIGHTGREEN_EX +"Developed By @Kunal"+Fore.LIGHTRED_EX +"========\n"+ Style.RESET_ALL)
    
    if status == "ocr":
        print(Fore.LIGHTCYAN_EX + "[OCR]"+ Fore.YELLOW + " Generating OCR Document" + Style.RESET_ALL)
   
        
        
# function for recreating output directory
def initialize_output_dir(dir_path):
    # Remove the directory if it exists
    if os.path.exists(dir_path):
        shutil.rmtree(dir_path)
    # Create a new result directory
    os.makedirs(dir_path)
    return dir_path
    

# function to perform OCR on images
def generate_ocr_document(folder_path,output_dir):
    print_function("ocr")
    
    # Create a new Word document
    document = Document()

    # Get the list of files in the folder
    files = sorted([os.path.join(folder_path, f) for f in os.listdir(folder_path) if f.lower().endswith('.png')])

    with progressbar.ProgressBar(max_value=len(files)) as bar:
        p = 0
        for index, file_name in enumerate(files):
            # Check if the file is an image
            if file_name.endswith(('.png')):
                file_path = os.path.join(folder_path, file_name)

                # Open the image
                image = Image.open(file_path)

                # Perform OCR using pytesseract
                ocr_text = pytesseract.image_to_string(image)

                # Add the image
                if (len(files) - p) <= 20: 
                    document.add_picture(file_path, width=Inches(3.9))
                else:
                    document.add_picture(file_path, width=Inches(7))
                # New Page Break
                document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                
                # Add the OCR result as text in the body of the section
                document.add_paragraph(ocr_text)
                
                # New Page Break
                document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            p+=1
            bar.update(p)

    # Save the Word document
    output_file = os.path.join(output_dir, 'ocr_result.docx')
    document.save(output_file)
    print("")
    
    
if __name__ == "__main__":
    # Load environment variables from .env file
    load_dotenv()
    
    pytesseract.pytesseract.tesseract_cmd = os.getenv("TESERRACT_PATH")
    
    print_function("intro")
    
    # Creating output directory
    input_dir = os.getenv("INPUT_IMAGE_DIR")
    output_dir = initialize_output_dir(os.getenv("OUTPUT_DIR"))
    
    # generating ocr document
    generate_ocr_document(input_dir,output_dir)
    
    
    print(Fore.LIGHTRED_EX + "=============================\n"+Style.BRIGHT + Fore.YELLOW +"            ENJOY            \n"+ Style.RESET_ALL + Fore.LIGHTRED_EX +"=============================\n"+ Style.RESET_ALL)
    input("Press 'Enter' To Close This Window")
